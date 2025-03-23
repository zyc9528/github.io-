import docx
import os
import re
import json

def extract_docx_content(file_path):
    """提取Word文档的内容"""
    try:
        doc = docx.Document(file_path)
        content = {
            'title': '',
            'paragraphs': [],
            'tables': [],
            'images': []
        }
        
        # 假设第一段是标题
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            content['title'] = doc.paragraphs[0].text.strip()
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                content['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name
                })
        
        # 提取表格
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append({
                'id': i,
                'data': table_data
            })
        
        # 图片信息（仅记录有图片的事实，无法直接提取图片）
        content['has_images'] = len(doc.inline_shapes) > 0
        
        return content
    except Exception as e:
        return {'error': str(e)}

def main():
    file_path = '乌市房地产市场 2025年深度研究报告（付费版）.docx'
    if not os.path.exists(file_path):
        print(f"文件不存在: {file_path}")
        return
    
    content = extract_docx_content(file_path)
    
    # 将内容保存为JSON文件，以便后续处理
    with open('report_content.json', 'w', encoding='utf-8') as f:
        json.dump(content, f, ensure_ascii=False, indent=2)
    
    print("内容已提取并保存到 report_content.json")

if __name__ == "__main__":
    main()